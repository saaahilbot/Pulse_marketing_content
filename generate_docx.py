from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x0d, 0x2b, 0x4e)
BLUE  = RGBColor(0x1a, 0x4f, 0x8a)
GREEN = RGBColor(0x1a, 0x7a, 0x4a)
MUTED = RGBColor(0x5a, 0x6a, 0x7e)
WHITE = RGBColor(0xff, 0xff, 0xff)

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── helpers ──────────────────────────────────────────────────────────────────

def cell_bg(cell, hex6):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    pr.append(shd)

def cell_border_left(cell, color_hex, sz=18):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(sz))
    left.set(qn('w:color'), color_hex)
    borders.append(left)
    pr.append(borders)

def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def para(text='', bold=False, size=10, color=None, space_before=0, space_after=4,
         align=WD_ALIGN_PARAGRAPH.LEFT, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p

def section_banner(text):
    t = doc.add_table(rows=1, cols=1)
    no_table_borders(t)
    c = t.cell(0,0)
    cell_bg(c, '0d2b4e')
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def initiative_row(number, title):
    t = doc.add_table(rows=1, cols=1)
    no_table_borders(t)
    c = t.cell(0,0)
    cell_bg(c, 'dce8f5')
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(10)
    r1 = p.add_run(f'INITIATIVE {number}    ')
    r1.bold = True; r1.font.size = Pt(8); r1.font.color.rgb = NAVY
    r2 = p.add_run(title)
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def bullet(text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Pt(18 if not sub else 28)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(9.5)

def outcome_table(rows):
    """Green-left-border box with month → outcome pairs."""
    t = doc.add_table(rows=len(rows)+1, cols=2)
    no_table_borders(t)
    # header
    c0 = t.cell(0,0); c1 = t.cell(0,1)
    cell_bg(c0, 'e6f4ed'); cell_bg(c1, 'e6f4ed')
    cell_border_left(c0, '1a7a4a', sz=18)
    r = c0.paragraphs[0].add_run('Expected outcomes')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GREEN
    # data rows
    for i, (month, outcome) in enumerate(rows, 1):
        cm = t.cell(i,0); co = t.cell(i,1)
        cell_bg(cm, 'f2faf5'); cell_bg(co, 'f2faf5')
        cell_border_left(cm, '1a7a4a', sz=6)
        rm = cm.paragraphs[0].add_run(month)
        rm.bold = True; rm.font.size = Pt(9)
        ro = co.paragraphs[0].add_run(outcome)
        ro.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK
# ─────────────────────────────────────────────────────────────────────────────
para('Vantage Pulse — Marketing Plan', bold=True, size=20, color=NAVY,
     space_before=0, space_after=2)
para('April 2026  ·  2-person team  ·  India + GCC focus', size=9, color=MUTED,
     space_before=0, space_after=10)

# ─────────────────────────────────────────────────────────────────────────────
# CONTEXT + FOCUS
# ─────────────────────────────────────────────────────────────────────────────
section_banner('Context')

# 2-col context table
ct = doc.add_table(rows=1, cols=2)
no_table_borders(ct)
c_left  = ct.cell(0,0)
c_right = ct.cell(0,1)
cell_bg(c_left,  'f0f4f9')
cell_bg(c_right, 'f0f4f9')

def fill_cell(cell, heading, lines):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    rh = p.add_run(heading + '\n')
    rh.bold = True; rh.font.size = Pt(9.5); rh.font.color.rgb = BLUE
    for line in lines:
        rn = p.add_run('• ' + line + '\n')
        rn.font.size = Pt(9)

fill_cell(c_left, 'Where we are', [
    'Leads come from Google Ads and 43 organic blogs',
    'Blogs convert — but the rate is low: no downloads, no content upgrades, no second path for buyers not ready to book a demo',
    'No G2 profile, no comparison pages, no product demo buyers can self-serve',
])
fill_cell(c_right, 'What we are building toward', [
    'Six initiatives a 2-person team can run independently',
    'Primary: India (IT services, BFSI, manufacturing, pharma) — 1,000–10,000 employee companies',
    'Secondary: Middle East / GCC (hospitality, financial services, logistics)',
    'USA captured passively via SEO — not actively targeted in Year 1',
])

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────────────────────────────────────
# COMPETITIVE CONTEXT
# ─────────────────────────────────────────────────────────────────────────────
section_banner('Competitive Landscape — Where the Gaps Are')

comp = doc.add_table(rows=5, cols=3)
comp.style = 'Table Grid'
for i, h in enumerate(['Competitor', 'Key weakness (from G2 reviews)', 'Our angle']):
    c = comp.cell(0,i)
    cell_bg(c, '1a4f8a')
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

rows_c = [
    ('Culture Amp  (1,543 reviews)',
     'No India office · ~$84K–$120K/year at 1,000 employees · anonymity concerns from reviewers',
     'Local support, genuinely anonymous, fraction of the cost at scale'),
    ('InFeedo / Amber  (81 reviews)',
     'CSV-based survey creation · G2 reviewers: "Creating reports for pulse surveys is not available"',
     'We are India-built too — but 4-click creation, AI-generated reports on every survey'),
    ('Lattice  (4,062 reviews)',
     'Pulse surveys are a secondary feature inside a performance platform',
     'We do one thing well. A dedicated listening tool, not a side tab'),
    ('15Five  (1,842 reviews)',
     'Built for small US teams · limited enterprise HRIS integrations',
     'Enterprise-grade: Workday, SAP, Oracle HCM, BambooHR out of the box'),
]
for ri, (comp_name, weakness, angle) in enumerate(rows_c, 1):
    for ci, txt in enumerate([comp_name, weakness, angle]):
        r = comp.cell(ri,ci).paragraphs[0].add_run(txt)
        r.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Key differentiator callout
diff_t = doc.add_table(rows=1, cols=1)
no_table_borders(diff_t)
dc = diff_t.cell(0,0)
cell_bg(dc, 'ede9fe')
p = dc.paragraphs[0]
p.paragraph_format.space_before = Pt(5)
p.paragraph_format.space_after  = Pt(5)
p.paragraph_format.left_indent  = Pt(10)
p.paragraph_format.right_indent = Pt(10)
r1 = p.add_run('Key differentiator no competitor has:  ')
r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = RGBColor(0x5b,0x21,0xb6)
r2 = p.add_run('AI-assisted anonymous conversation thread — employee leaves feedback, HR opens the thread, AI suggests three empathetic responses at 94% confidence. Not Culture Amp, not InFeedo, not Lattice offers this. Lead with it on every comparison page and demo.')
r2.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────────────────────────────────────
# THE 6 INITIATIVES
# ─────────────────────────────────────────────────────────────────────────────
section_banner('The 6 Initiatives')

# ── 1 ──
initiative_row('1', 'Establish G2 Presence')
para('Create Vantage Pulse profiles on G2 and Capterra, run a structured review collection campaign targeting existing clients at 1,000+ employee companies.', size=9.5, space_before=0, space_after=4)

para('Tactics', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
for b in [
    'Claim G2 + Capterra profiles — tag all 3 categories: Employee Engagement, Pulse Survey, Employee Feedback Software',
    'Upload product screenshots, compliance badges (ISO 27001, SOC 2, GDPR, HIPAA), and HRIS integrations',
    'Send 10 personalised review requests/week to existing Pulse clients for the first 4 weeks; 5/week as ongoing drip',
    'Incentivise via Vantage Perks. Ask for honest reviews — G2 flags patterns that look artificial',
]:
    bullet(b)

outcome_table([
    ('Month 1', '15–20 reviews. Profile live and appearing in brand searches.'),
    ('Month 2', '40–60 reviews. Category grid threshold (20+) crossed. Pulse appears in comparison grids. 2–4 passive inbound demo requests/month.'),
    ('Month 3', '80–100 reviews. High Performer badge eligibility. 4–7 inbound demo requests/month from G2 traffic.'),
    ('Month 6', '120–160 reviews. G2 a consistent source of 6–9 qualified demo requests/month with no additional effort.'),
])

# ── 2 ──
initiative_row('2', 'Build 6 Comparison Pages (1 per month)')
para('Publish bottom-of-funnel SEO pages targeting buyers actively choosing between tools. Currently zero comparison pages exist — every buyer searching "Vantage Pulse vs Culture Amp" lands on a competitor.', size=9.5, space_before=0, space_after=4)

para('Page schedule', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
pages = [
    ('Month 1', '"Vantage Pulse vs Culture Amp"', "India gap, cost at scale, anonymity concerns"),
    ('Month 2', '"Best Culture Amp alternatives for India"', "Uncontested search — no competitor targets India specifically"),
    ('Month 3', '"Vantage Pulse vs InFeedo"', "CSV creation weakness, missing pulse reports"),
    ('Month 4', '"Best employee pulse survey software 2026"', "Own the roundup; lead with AI response assistant"),
    ('Month 5', '"Employee engagement surveys for 1,000+ employee companies"', "Use-case page for the ICP"),
    ('Month 6', '"Vantage Pulse vs Lattice"', "Dedicated tool vs side feature in a performance platform"),
]
for month, title, angle in pages:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Pt(18)
    r1 = p.add_run(f'{month}  {title} — ')
    r1.bold = True; r1.font.size = Pt(9.5)
    r2 = p.add_run(angle)
    r2.font.size = Pt(9.5)

outcome_table([
    ('Month 1–3', 'Pages indexed. Internal links from existing blogs build authority. Minimal traffic — SEO has a 3–5 month lag.'),
    ('Month 4–5', 'First pages begin ranking. 100–300 visits/month per page — decision-stage traffic with the highest intent.'),
    ('Month 6',   '3–4 pages on page 1. Combined 400–1,200 visits/month → 8–36 incremental demo requests/month. India-specific pages rank faster — no competition.'),
])

# ── 3 ──
initiative_row('3', 'Build an Interactive Demo')
para('A guided click-through of the Pulse UI (Navattic or Storylane, ~$400/month, 1-day build). Embedded on the product page and every comparison page. Converts visitors who want to evaluate before committing to a sales call.', size=9.5, space_before=0, space_after=4)

para('6-screen flow', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
for b in [
    'Dashboard — Engagement Index 65, eNPS 67, Participation 36% flagged low',
    'Heatmap — Technical Support, Engagement: score of 04. The urgency screen.',
    'AI response assistant — HR replies to anonymous feedback; AI suggests 3 empathetic drafts at 94% confidence',
    'Lifecycle timeline — automated triggers at Day 7/30/90, Year 1/3/5',
    'HRIS integrations — Workday, SAP, Oracle HCM, BambooHR, Okta, Azure AD',
    'CTA — "Book a 15-min demo" + "See a sample engagement report"',
]    :
    bullet(b)

outcome_table([
    ('Month 1',  'Demo live. Baseline data collecting.'),
    ('Month 2',  'Drop-off data available. Fix weakest annotation. Target completion rate: >25%.'),
    ('Month 3+', 'At 1,500 product page visits/month: 375–600 engage with demo → 55–150 call requests/month. Largest single conversion lift from a one-time build.'),
])

# ── 4 ──
initiative_row('4', 'Two Gated Lead Assets')
para('Build two downloadable assets that create a second conversion path: visit → download → 2-email nurture sequence → demo. Captures buyers who are researching but not ready to talk to sales.', size=9.5, space_before=0, space_after=4)

para('The two assets', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
for b in [
    'Employee Engagement Benchmark Report (6–8 page PDF) — engagement scores by industry (IT services, BFSI, manufacturing, hospitality), eNPS benchmarks for India + GCC, participation rate norms by company size. Built from anonymised Pulse data. Gate: work email + company size.',
    'eNPS and Disengagement Cost Calculator (Typeform) — HR inputs headcount, participation rate, eNPS estimate. Output: benchmark comparison + estimated annual cost of disengagement. "Estimated disengagement cost at your headcount: $1.8M/year" gets shared internally. Gate: work email.',
]:
    bullet(b)
bullet('Post-download: 2-email automated sequence. Email 1 delivers the asset + interactive demo link. Email 2 (Day 4): one industry-specific insight, soft CTA.', sub=False)
bullet('Anyone who selects 1,000+ employees at download is flagged as a priority lead for sales.', sub=False)

outcome_table([
    ('Month 2', 'Both assets live on gated landing pages. Embedded on comparison pages and G2 profile.'),
    ('Month 3', '30–80 downloads/month. 10–25 qualified leads (1,000+ filter) captured/month.'),
    ('Month 6', '60–150 downloads/month. 3–10 incremental demo requests/month from this channel alone.'),
])

# ── 5 ──
initiative_row('5', 'Content Upgrades Across the 43 Blogs')
para('Blogs already generate leads — but readers have no second action besides "Book a Demo." Add relevant downloadable assets inline within each blog cluster to increase conversion without writing new content.', size=9.5, space_before=0, space_after=4)

para('6 downloads mapped to blog clusters', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
upgrades = doc.add_table(rows=7, cols=3)
upgrades.style = 'Table Grid'
for i, h in enumerate(['Blog cluster', 'Download', 'Format']):
    c = upgrades.cell(0,i)
    cell_bg(c, 'dce8f5')
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLUE
urows = [
    ('Survey strategy, fatigue, timing (9 posts)',       'Survey Launch Checklist — 12 steps before sending',    'PDF'),
    ('All 16 question posts',                            '47 Validated Pulse Survey Questions (editable)',        'Word'),
    ('Participation + anonymous surveys',                '3 Survey Communication Email Templates',               'Word'),
    ('eNPS + engagement + satisfaction blogs',           'Benchmark Report + eNPS Calculator (Initiative 4)',    'PDF / Typeform'),
    ('Feedback + action planning',                       'Post-Survey Action Planning Template',                  'Google Sheet'),
    ('Lifecycle + retention + culture + climate',        'Lifecycle Survey Calendar + Retention Risk Assessment', 'Google Sheet / PDF'),
]
for ri, (cluster, dl, fmt) in enumerate(urows, 1):
    for ci, txt in enumerate([cluster, dl, fmt]):
        r = upgrades.cell(ri,ci).paragraphs[0].add_run(txt)
        r.font.size = Pt(9)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

outcome_table([
    ('Month 2', 'First 4 downloads live. Placed inline on 30+ relevant blog posts.'),
    ('Month 3', '50–150 downloads/month from existing blog traffic. 15–50 qualified leads/month — same traffic, more of it converted.'),
    ('Month 6', 'Blog lead conversion rate measurably higher than baseline. 8–30 incremental demo requests/month from content-driven leads.'),
])

# ── 6 ──
initiative_row('6', 'LinkedIn Organic Presence')
para('Post 2–3x/week on the company page. LinkedIn is not the primary pipeline driver — its job is brand recognition. A buyer who has seen Pulse posts before visiting the product page converts at a higher rate.', size=9.5, space_before=0, space_after=4)

para('Three post formats (rotating)', bold=True, size=9.5, color=BLUE, space_before=2, space_after=2)
for b in [
    'Data post — lead with a real number (e.g. "One department had an engagement score of 04. No manager flagged it. The heatmap caught it."). No product pitch.',
    'Product visual post — one screenshot with a factual caption. The AI response assistant screenshot is the most powerful.',
    'Opinion post — one short take designed to generate comments. No product mention.',
]:
    bullet(b)
bullet('Monthly Loom video: 90-second screen recording of the live dashboard, no editing, posted as native LinkedIn video. Reaches 3–5x more people than any other post format. 30 minutes total effort.', sub=False)

outcome_table([
    ('Month 1–2', '1,000–8,000 impressions/month. Build the habit — audience does not exist yet.'),
    ('Month 3',   '15,000–30,000 impressions/month. Engagement rate >2%. 1–3 direct enquiries/month. 20–50 benchmark report downloads/month from LinkedIn.'),
    ('Month 6',   '40,000–80,000 impressions/month. 4–7 inbound enquiries/month. Higher sales reply rate from prospects who already follow the page.'),
])

# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
section_banner('Expected Pipeline — Incremental Demo Requests Above Current Baseline')

pt = doc.add_table(rows=7, cols=4)
pt.style = 'Table Grid'
for i, h in enumerate(['Channel', 'Month 2', 'Month 3', 'Month 6']):
    c = pt.cell(0,i)
    cell_bg(c, '1a4f8a')
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE

prows = [
    ('G2 / Capterra — passive inbound',        '0–2',   '4–7',   '6–9'),
    ('Comparison pages — organic search',      '0',     '0–2',   '8–36'),
    ('Interactive demo — product page lift',   '15–30', '20–40', '30–60'),
    ('Gated assets + blog downloads',          '0–2',   '3–8',   '8–30'),
    ('LinkedIn — direct enquiries',            '0',     '1–3',   '4–7'),
    ('Total incremental / month',              '15–34', '28–60', '56–142'),
]
for ri, (ch, m2, m3, m6) in enumerate(prows, 1):
    is_total = ri == 6
    for ci, txt in enumerate([ch, m2, m3, m6]):
        c = pt.cell(ri, ci)
        if is_total:
            cell_bg(c, 'dce8f5')
        r = c.paragraphs[0].add_run(txt)
        r.bold = is_total
        r.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
para('All figures are incremental above the existing Google Ads + blog baseline. SEO comparison pages begin generating traffic at Month 4–5. Interactive demo projected from 1,500 product page visits/month.', size=8.5, color=MUTED, space_before=0, space_after=8)

# ─────────────────────────────────────────────────────────────────────────────
# BUDGET + METRICS
# ─────────────────────────────────────────────────────────────────────────────
bt = doc.add_table(rows=1, cols=2)
no_table_borders(bt)
bl = bt.cell(0,0)
br_cell = bt.cell(0,1)

# Budget
cell_bg(bl, 'f0f4f9')
pb = bl.paragraphs[0]
pb.paragraph_format.space_before = Pt(6)
pb.paragraph_format.space_after  = Pt(4)
pb.paragraph_format.left_indent  = Pt(8)
rh = pb.add_run('Monthly Budget\n')
rh.bold = True; rh.font.size = Pt(10); rh.font.color.rgb = BLUE
budget_lines = [
    ('Interactive demo (Navattic / Storylane)', '~$400'),
    ('G2 review incentives (Vantage Perks)',    '~$200'),
    ('Calculator hosting (Typeform)',            '~$50'),
    ('Email automation (HubSpot / Mailchimp)',   '$0–50'),
    ('Content (internal, Claude skills)',         '$0'),
    ('Total',                                    '~$650–700'),
]
for item, cost in budget_lines:
    is_total = item == 'Total'
    r = pb.add_run(('━' * 28 + '\n') if is_total else '')
    r.font.size = Pt(7)
    line = pb.add_run(f'{item:<38}{cost}\n')
    line.font.size = Pt(9)
    if is_total:
        line.bold = True

# Metrics
cell_bg(br_cell, 'f0f4f9')
pm = br_cell.paragraphs[0]
pm.paragraph_format.space_before = Pt(6)
pm.paragraph_format.space_after  = Pt(4)
pm.paragraph_format.left_indent  = Pt(8)
rh2 = pm.add_run('Key Metrics and Red Flags\n')
rh2.bold = True; rh2.font.size = Pt(10); rh2.font.color.rgb = BLUE
metrics = [
    ('G2 reviews — end of Month 3',   'Below 80 → review process is broken'),
    ('Demo completion rate',           'Below 20% → Step 2 annotation needs rewriting'),
    ('Blog download rate',             'Below 20/month by Month 3 → fix placement or copy'),
    ('Comparison page ranking',        'No movement by Month 4 → add internal links'),
    ('LinkedIn engagement rate',       'Below 1.5% → switch to data + visual posts'),
]
for metric, flag in metrics:
    rm = pm.add_run(f'{metric}\n')
    rm.bold = True; rm.font.size = Pt(9)
    rf = pm.add_run(f'  ↳ {flag}\n')
    rf.font.size = Pt(8.5); rf.font.color.rgb = MUTED

pm.add_run('\n')
rns = pm.add_run('North star: ')
rns.bold = True; rns.font.size = Pt(9)
rnsv = pm.add_run('demo requests/month from non-Ads channels. Month 3: 28+. Month 6: 56+.')
rnsv.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
para('Google Ads is separate (already running). LinkedIn Paid: activate at Month 4–5 once 3 comparison pages and 1 new case study are live. Start at $1,500–2,000/month, targeting HR Directors + CHROs in India + GCC.',
     size=8.5, color=MUTED, space_before=0, space_after=4)

# ─────────────────────────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────────────────────────
para('Vantage Pulse Marketing Plan · April 2026 · 2-person team · India + GCC enterprise focus · Competitive data from G2 review analysis · India HR tech market: IMARC Research 2024',
     size=7.5, color=MUTED, space_before=8, space_after=0)

doc.save('/Users/sahil/Desktop/marketing/pulse-marketing-plan.docx')
print('Done')
